"""Word 文档加载器"""
from pathlib import Path
from typing import List, Union
from langchain_core.documents import Document


class DocxLoader:
    """Word (.docx/.doc) 文档加载器"""

    def load(self, file_path: Union[str, Path]) -> List[Document]:
        """加载单个 Word 文档"""
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == ".docx":
            return self._load_docx(file_path)
        elif suffix == ".doc":
            return self._load_doc(file_path)
        else:
            return []

    def _load_docx(self, file_path: Path) -> List[Document]:
        """加载 .docx 文件"""
        import docx

        doc = docx.Document(str(file_path))
        # 提取所有段落文本
        content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        # 提取表格内容
        tables_content = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    tables_content.append(row_text)

        if tables_content:
            content += "\n\n【表格内容】\n" + "\n".join(tables_content)

        metadata = {
            "source": str(file_path),
            "title": file_path.stem,
            "file_type": ".docx"
        }

        return [Document(page_content=content, metadata=metadata)]

    def _load_doc(self, file_path: Path) -> List[Document]:
        """加载 .doc 文件（旧格式）"""
        content = None

        # 方案1：使用 pywin32（仅 Windows，推荐）
        try:
            import win32com.client
            import pythoncom

            pythoncom.CoInitialize()
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(file_path.absolute()))

                # 提取段落
                paragraphs = [p.Range.Text.strip() for p in doc.Paragraphs if p.Range.Text.strip()]

                # 提取表格
                tables_content = []
                for table in doc.Tables:
                    for row in table.Rows:
                        row_text = " | ".join([cell.Range.Text.strip() for cell in row.Cells])
                        if row_text.strip():
                            tables_content.append(row_text)

                doc.Close(False)
                word.Quit()

                # 合并段落和表格
                content = "\n".join(paragraphs)
                if tables_content:
                    content += "\n\n【设备清单表格】\n" + "\n".join(tables_content)
            finally:
                pythoncom.CoUninitialize()
        except Exception as e1:
            # 方案2：使用 doc2txt
            try:
                import doc2txt
                content = doc2txt.extract_text(str(file_path), optimize_format=True)
            except Exception as e2:
                # 方案3：直接读取二进制，手动解析（最后手段）
                try:
                    content = self._extract_doc_text_fallback(file_path)
                except Exception as e3:
                    print(f"[FAIL] .doc 解析失败: pywin32={e1}, doc2txt={e2}, fallback={e3}")
                    return []

        if not content or not content.strip():
            return []

        metadata = {
            "source": str(file_path),
            "title": file_path.stem,
            "file_type": ".doc"
        }

        return [Document(page_content=content.strip(), metadata=metadata)]

    def _extract_doc_text_fallback(self, file_path: Path) -> str:
        """使用 antiword 直接读取 .doc 文件"""
        import subprocess
        import tempfile

        # 尝试不同的 antiword 路径
        possible_paths = [
            "antiword",
            r"C:\Program Files\antiword\antiword.exe",
            r"C:\antiword\antiword.exe",
        ]

        encodings = ["gb2312", "gbk", "gb18030", "utf-8"]

        for antiword_path in possible_paths:
            for encoding in encodings:
                try:
                    result = subprocess.run(
                        [antiword_path, "-f", str(file_path)],
                        capture_output=True,
                        timeout=30
                    )
                    if result.returncode == 0:
                        try:
                            return result.stdout.decode(encoding)
                        except:
                            continue
                except:
                    continue

        raise RuntimeError("antiword not found or failed")

    def load_directory(self, directory: Union[str, Path]) -> List[Document]:
        """加载目录下所有 Word 文档"""
        directory = Path(directory)
        documents = []

        # 支持 .docx 和 .doc 格式
        for doc_file in directory.rglob("*.docx"):
            try:
                docs = self.load(doc_file)
                documents.extend(docs)
                print(f"[OK] 加载: {doc_file.name}")
            except Exception as e:
                print(f"[FAIL] 加载失败 {doc_file.name}: {e}")

        for doc_file in directory.rglob("*.doc"):
            try:
                docs = self.load(doc_file)
                documents.extend(docs)
                print(f"[OK] 加载: {doc_file.name}")
            except Exception as e:
                print(f"[FAIL] 加载失败 {doc_file.name}: {e}")

        return documents
